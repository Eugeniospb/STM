"""
Фемида — Юридический ассистент ООО "СТМ"
Версия: 2.0
Дата: Январь 2026

ДОСТУП:
- Личка: ТОЛЬКО директор (@eugenio_spb)
- Группа "Наше производство": ТОЛЬКО директор с триггером "Фемида,"

МОДЕЛИ:
- Haiku (дешёвая): простые вопросы, справки, быстрые ответы
- Sonnet (дорогая): анализ документов, генерация договоров, сложные задачи

ВОЗМОЖНОСТИ:
- Обработка PDF и изображений (Vision API)
- Генерация DOCX на фирменном бланке СТМ с логотипом
"""

import os
import io
import re
import json
import base64
import logging
from datetime import datetime
from typing import Optional
from pathlib import Path

from telegram import Update, Chat, Message, Document, PhotoSize
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    ContextTypes, filters
)
from telegram.constants import ParseMode, ChatAction

import anthropic
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ==================== ЛОГИРОВАНИЕ ====================

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)


# ==================== КОНФИГУРАЦИЯ ====================

# Токены
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

# AI модели (tiered) - ИСПРАВЛЕННЫЕ НАЗВАНИЯ
MODEL_CHEAP = "claude-3-haiku-20240307"      # Для простых запросов
MODEL_EXPENSIVE = "claude-sonnet-4-20250514"  # Для документов и анализа

# Лимиты
MAX_TOKENS_CHEAP = 2048
MAX_TOKENS_EXPENSIVE = 4096

# ДОСТУП: Только директор
DIRECTOR_USERNAME = "eugenio_spb"
DIRECTOR_ID = 1676748258

# Группа "Наше производство"
GROUP_ID = int(os.getenv("GROUP_ID", "-1003639268911"))

# Триггеры (регистронезависимо)
TRIGGERS = ["фемида,", "феми,", "фемида ", "феми "]

# Путь к ассетам
ASSETS_DIR = Path(__file__).parent / "assets"
LOGO_PATH = ASSETS_DIR / "logo.png"


# ==================== РЕКВИЗИТЫ СТМ ====================

COMPANY = {
    "full_name": "Общество с ограниченной ответственностью «СТМ»",
    "short_name": "ООО «СТМ»",
    "inn": "7813568956",
    "kpp": "781401001",
    "ogrn": "1137847312866",
    "address": "197375, Санкт-Петербург, ул. Маршала Новикова д.42, Литер А, Помещение ПИБ №1-Н-113",
    "bank": "АО «ТИНЬКОФФ БАНК»",
    "bik": "044525974",
    "rs": "40702810810000134609",
    "ks": "30101810145250000974",
    "director": "Тихонов Евгений Викторович",
    "director_short": "Тихонов Е.В.",
    "director_position": "Генеральный директор",
    "phone": "+7 812 603 78 71",
    "email": "stm.laser@gmail.com",
}

IP_TIKHONOV = {
    "full_name": "Индивидуальный предприниматель Тихонов Александр Викторович",
    "short_name": "ИП Тихонов А.В.",
    "inn": "781428127765",
    "ogrnip": "319784700268498",
    "address": "197375, Санкт-Петербург, ул. Репищева д.17, корп.1, кв.28",
    "bank": "АО «ТИНЬКОФФ БАНК»",
    "bik": "044525974",
    "rs": "40802810400001208048",
    "ks": "30101810145250000974",
}


# ==================== КЛИЕНТ ANTHROPIC ====================

client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)


# ==================== ОПРЕДЕЛЕНИЕ СЛОЖНОСТИ ЗАПРОСА ====================

EXPENSIVE_PATTERNS = [
    r"(составь|напиши|подготовь|создай|сделай).*(договор|письмо|претензи|приказ|иск|заявлени|акт|счёт|счет)",
    r"(договор|письмо|претензи|иск).*(на имя|в адрес|для)",
    r"(проанализируй|проверь|изучи|оцени).*(договор|документ|контракт|соглашени)",
    r"(что не так|ошибки|риски|проблемы).*(договор|документ|контракт)",
    r"(как (подать|составить|оформить|написать)).*(иск|претензи|жалоб)",
    r"(взыскать|вернуть).*(деньги|долг|задолженность)",
    r"(разработай|предложи).*(стратеги|план|схем)",
]


def is_expensive_request(text: str, has_file: bool = False) -> bool:
    """Определяет, нужна ли дорогая модель"""
    # Файлы всегда через дорогую модель (Vision)
    if has_file:
        return True
    
    text_lower = text.lower()
    
    for pattern in EXPENSIVE_PATTERNS:
        if re.search(pattern, text_lower):
            return True
    
    if len(text) > 500:
        return True
    
    return False


def get_model_for_request(text: str, has_file: bool = False) -> tuple:
    """Возвращает модель и лимит токенов"""
    if is_expensive_request(text, has_file):
        return MODEL_EXPENSIVE, MAX_TOKENS_EXPENSIVE
    return MODEL_CHEAP, MAX_TOKENS_CHEAP


# ==================== ПРОВЕРКА ДОСТУПА ====================

def is_director(user_id: int, username: str = None) -> bool:
    """Проверяет, директор ли это"""
    if user_id == DIRECTOR_ID:
        return True
    if username and username.lower() == DIRECTOR_USERNAME.lower():
        return True
    return False


def has_trigger(text: str) -> tuple:
    """Проверяет наличие триггера"""
    text_lower = text.lower()
    for trigger in TRIGGERS:
        if text_lower.startswith(trigger):
            return True, text[len(trigger):].strip()
    return False, text


# ==================== ОБРАБОТКА ФАЙЛОВ ====================

async def download_file(bot, file_id: str) -> bytes:
    """Скачивает файл из Telegram"""
    file = await bot.get_file(file_id)
    buffer = io.BytesIO()
    await file.download_to_memory(buffer)
    buffer.seek(0)
    return buffer.read()


async def process_document(bot, document: Document) -> tuple:
    """
    Обрабатывает документ, возвращает (base64_data, media_type)
    """
    file_name = document.file_name or "file"
    mime_type = document.mime_type or "application/octet-stream"
    
    file_data = await download_file(bot, document.file_id)
    base64_data = base64.standard_b64encode(file_data).decode("utf-8")
    
    # Определяем media_type для Claude
    if mime_type == "application/pdf":
        return base64_data, "application/pdf"
    elif mime_type.startswith("image/"):
        return base64_data, mime_type
    else:
        # Для других файлов пытаемся как текст
        try:
            text_content = file_data.decode("utf-8")
            return text_content, "text"
        except:
            return base64_data, mime_type
    
    return None, None


async def process_photo(bot, photo: PhotoSize) -> tuple:
    """Обрабатывает фото"""
    file_data = await download_file(bot, photo.file_id)
    base64_data = base64.standard_b64encode(file_data).decode("utf-8")
    return base64_data, "image/jpeg"


# ==================== ГЕНЕРАЦИЯ ОТВЕТА ====================

def get_current_date_ru() -> str:
    """Текущая дата в русском формате"""
    months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    now = datetime.now()
    return f"{now.day} {months[now.month]} {now.year} г."


def build_system_prompt() -> str:
    """Системный промпт для Фемиды"""
    return f"""Ты — юридический ассистент "Фемида" компании {COMPANY['short_name']}.

ТВОИ ЗАДАЧИ:
- Составление юридических документов (договоры, письма, претензии, приказы, иски)
- Анализ договоров и выявление рисков
- Консультации по правовым вопросам (ГК РФ, ТК РФ, НК РФ, АПК РФ)
- Анализ присланных документов (PDF, изображения)

СТИЛЬ:
- Официально-деловой, но дружелюбный
- Обращайся к директору на "вы" или по имени "Евгений"
- Ссылайся на конкретные статьи законов
- Будь практичным — давай готовые решения

РЕКВИЗИТЫ ООО «СТМ»:
ИНН: {COMPANY['inn']}
КПП: {COMPANY['kpp']}
ОГРН: {COMPANY['ogrn']}
Адрес: {COMPANY['address']}
Р/с: {COMPANY['rs']}
Банк: {COMPANY['bank']}
БИК: {COMPANY['bik']}
Директор: {COMPANY['director']}

РЕКВИЗИТЫ ИП Тихонов А.В.:
ИНН: {IP_TIKHONOV['inn']}
ОГРНИП: {IP_TIKHONOV['ogrnip']}
Адрес: {IP_TIKHONOV['address']}
Р/с: {IP_TIKHONOV['rs']}

СЕГОДНЯ: {get_current_date_ru()}

ВАЖНО:
- При составлении документов используй правильные реквизиты
- Если просят "от ИП" — используй ИП Тихонов А.В.
- По умолчанию документы от ООО СТМ
- Если прислали файл — анализируй его содержимое
"""


async def generate_response(text: str, file_data: tuple = None) -> tuple:
    """
    Генерирует ответ через Claude.
    file_data: (base64_data, media_type) или None
    Возвращает (ответ, использованная_модель)
    """
    has_file = file_data is not None
    model, max_tokens = get_model_for_request(text, has_file)
    
    logger.info(f"Запрос: '{text[:50]}...' → модель: {model}, файл: {has_file}")
    
    try:
        # Формируем контент сообщения
        if file_data and file_data[0]:
            base64_data, media_type = file_data
            
            if media_type == "text":
                # Текстовый файл — добавляем в промпт
                content = [
                    {"type": "text", "text": f"Содержимое файла:\n\n{base64_data}\n\nЗапрос: {text}"}
                ]
            elif media_type == "application/pdf":
                # PDF через document type
                content = [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": base64_data
                        }
                    },
                    {"type": "text", "text": text if text else "Проанализируй этот документ. Что это и о чём он?"}
                ]
            else:
                # Изображение
                content = [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": base64_data
                        }
                    },
                    {"type": "text", "text": text if text else "Что изображено на этом документе? Проанализируй."}
                ]
        else:
            content = [{"type": "text", "text": text}]
        
        message = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=build_system_prompt(),
            messages=[{"role": "user", "content": content}]
        )
        
        response_text = message.content[0].text
        
        input_tokens = message.usage.input_tokens
        output_tokens = message.usage.output_tokens
        logger.info(f"Токены: in={input_tokens}, out={output_tokens}, модель={model}")
        
        return response_text, model
        
    except anthropic.NotFoundError as e:
        logger.error(f"Модель не найдена: {e}")
        return f"⚠️ Ошибка: модель {model} недоступна. Обратитесь к администратору.", model
    except Exception as e:
        logger.error(f"Ошибка Claude API: {e}")
        return f"⚠️ Ошибка генерации ответа: {e}", model


# ==================== ГЕНЕРАЦИЯ DOCX НА БЛАНКЕ ====================

def create_docx_on_letterhead(content: str, title: str = "Документ") -> io.BytesIO:
    """
    Создаёт DOCX на фирменном бланке ООО СТМ
    """
    doc = DocxDocument()
    
    # ===== СТИЛИ =====
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(0)
    
    # ===== ПОЛЯ (ГОСТ Р 7.0.97-2016) =====
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)  # Для подшивки
        section.right_margin = Cm(1.5)
        section.header_distance = Cm(1)
    
    # ===== ШАПКА (HEADER) =====
    header = doc.sections[0].header
    
    # Таблица для шапки (логотип слева, текст справа)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    
    # Ширина колонок
    header_table.columns[0].width = Inches(1.2)
    header_table.columns[1].width = Inches(5.3)
    
    # Логотип (если есть)
    logo_cell = header_table.cell(0, 0)
    if LOGO_PATH.exists():
        logo_para = logo_cell.paragraphs[0]
        logo_run = logo_para.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Inches(1))
    
    # Текст шапки
    text_cell = header_table.cell(0, 1)
    
    # Название компании
    name_para = text_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run("ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ")
    name_run.font.name = 'Times New Roman'
    name_run.font.size = Pt(11)
    name_run.font.bold = True
    
    name_para2 = text_cell.add_paragraph()
    name_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run2 = name_para2.add_run("«СТМ»")
    name_run2.font.name = 'Times New Roman'
    name_run2.font.size = Pt(14)
    name_run2.font.bold = True
    name_run2.font.color.rgb = RGBColor(0, 112, 192)  # Синий
    
    # Реквизиты
    details_para = text_cell.add_paragraph()
    details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_text = f"Россия, {COMPANY['address']}"
    details_run = details_para.add_run(details_text)
    details_run.font.name = 'Times New Roman'
    details_run.font.size = Pt(8)
    
    details_para2 = text_cell.add_paragraph()
    details_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_run2 = details_para2.add_run(f"ИНН {COMPANY['inn']} · КПП {COMPANY['kpp']} · ОГРН {COMPANY['ogrn']}")
    details_run2.font.name = 'Times New Roman'
    details_run2.font.size = Pt(8)
    
    # Линия под шапкой
    line_para = header.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_para.add_run("─" * 85)
    line_run.font.size = Pt(8)
    line_run.font.color.rgb = RGBColor(0, 112, 192)
    
    # ===== ОСНОВНОЙ КОНТЕНТ =====
    # Пустая строка после шапки
    doc.add_paragraph()
    
    # Контент документа
    for para_text in content.split('\n'):
        if para_text.strip():
            p = doc.add_paragraph()
            
            # Определяем тип параграфа
            stripped = para_text.strip()
            
            # Заголовки (всё заглавными или ключевые слова)
            if stripped.isupper() or any(stripped.startswith(x) for x in 
                ['ДОГОВОР', 'ПРИКАЗ', 'ПРЕТЕНЗИЯ', 'ИСКОВОЕ', 'АКТ', 'ПИСЬМО', 'ЗАЯВЛЕНИЕ']):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(14)
            
            # Даты и номера (справа)
            elif stripped.startswith(('г.', 'от ', '«', '"')) and len(stripped) < 50:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(stripped)
            
            # Обычный текст
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Первая строка с отступом
                p.paragraph_format.first_line_indent = Cm(1.25)
                run = p.add_run(stripped)
            
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    # ===== ПОДПИСЬ =====
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_para = doc.add_paragraph()
    sig_para.add_run(f"{COMPANY['director_position']}")
    sig_para.add_run("                    ")
    sig_para.add_run("_____________")
    sig_para.add_run("    ")
    sig_para.add_run(f"{COMPANY['director_short']}")
    
    # М.П.
    mp_para = doc.add_paragraph()
    mp_para.add_run("                              М.П.")
    for run in mp_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    # ===== ФУТЕР =====
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"Тел: {COMPANY['phone']} · Email: {COMPANY['email']}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Сохраняем
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==================== ОБРАБОТЧИКИ TELEGRAM ====================

async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Команда /start"""
    user = update.effective_user
    
    if not is_director(user.id, user.username):
        await update.message.reply_text(
            "⚖️ Фемида — юридический ассистент ООО «СТМ».\n\n"
            "Бот доступен только для руководства компании."
        )
        return
    
    await update.message.reply_text(
        "⚖️ *Фемида* — юридический ассистент ООО «СТМ»\n\n"
        "Я помогу с:\n"
        "• Составлением договоров, писем, претензий\n"
        "• Анализом документов (PDF, фото)\n"
        "• Юридическими консультациями\n\n"
        "Просто напишите ваш вопрос или пришлите документ.\n\n"
        "_В группе используйте триггер: Фемида, ваш вопрос_",
        parse_mode=ParseMode.MARKDOWN
    )


async def cmd_help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Команда /help"""
    await update.message.reply_text(
        "⚖️ *Возможности Фемиды:*\n\n"
        "*Документы:*\n"
        "• «Составь письмо в адрес ООО Ромашка»\n"
        "• «Напиши претензию на возврат денег»\n"
        "• «Подготовь договор на изготовление вывески»\n\n"
        "*Анализ файлов:*\n"
        "• Пришлите PDF или фото + вопрос\n"
        "• «Что это за документ?»\n"
        "• «Проверь на риски»\n\n"
        "*Консультации:*\n"
        "• «Какой срок исковой давности?»\n"
        "• «Можно ли расторгнуть договор?»\n\n"
        "*Реквизиты:*\n"
        "• /requisites — реквизиты СТМ\n"
        "• «Реквизиты ИП»\n\n"
        "_Документы генерируются на фирменном бланке_",
        parse_mode=ParseMode.MARKDOWN
    )


async def cmd_requisites(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Команда /requisites"""
    text = f"""📋 *Реквизиты ООО «СТМ»*

*Полное наименование:*
{COMPANY['full_name']}

*ИНН:* `{COMPANY['inn']}`
*КПП:* `{COMPANY['kpp']}`
*ОГРН:* `{COMPANY['ogrn']}`

*Адрес:*
{COMPANY['address']}

*Банковские реквизиты:*
Банк: {COMPANY['bank']}
Р/с: `{COMPANY['rs']}`
К/с: `{COMPANY['ks']}`
БИК: `{COMPANY['bik']}`

*Контакты:*
Тел: {COMPANY['phone']}
Email: {COMPANY['email']}

*Генеральный директор:*
{COMPANY['director']}
"""
    await update.message.reply_text(text, parse_mode=ParseMode.MARKDOWN)


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Главный обработчик сообщений"""
    message = update.message
    if not message:
        return
    
    user = update.effective_user
    chat = update.effective_chat
    
    # Получаем текст и файлы
    text = message.text or message.caption or ""
    text = text.strip()
    
    # Проверяем наличие файлов
    file_data = None
    
    if message.document:
        file_data = await process_document(context.bot, message.document)
    elif message.photo:
        # Берём самое большое фото
        photo = message.photo[-1]
        file_data = await process_photo(context.bot, photo)
    
    # Если только файл без текста
    if file_data and not text:
        text = "Проанализируй этот документ. Что это и о чём он?"
    
    # Если нет ни текста, ни файла — игнорируем
    if not text and not file_data:
        return
    
    # ========== ЛИЧНЫЙ ЧАТ ==========
    if chat.type == Chat.PRIVATE:
        if not is_director(user.id, user.username):
            await message.reply_text("⚖️ Фемида доступна только руководству ООО «СТМ».")
            return
        
        await process_request(message, text, file_data, context)
        return
    
    # ========== ГРУППОВОЙ ЧАТ ==========
    if chat.type in [Chat.GROUP, Chat.SUPERGROUP]:
        if chat.id != GROUP_ID:
            return
        
        if not is_director(user.id, user.username):
            return
        
        has_trig, clean_text = has_trigger(text)
        if not has_trig:
            return
        
        await process_request(message, clean_text, file_data, context)
        return


async def process_request(message: Message, text: str, file_data: tuple, context: ContextTypes.DEFAULT_TYPE):
    """Обрабатывает запрос и отправляет ответ"""
    
    # Показываем "печатает..."
    await context.bot.send_chat_action(
        chat_id=message.chat_id,
        action=ChatAction.TYPING
    )
    
    # Быстрые команды на реквизиты
    text_lower = text.lower()
    if "реквизиты" in text_lower and not file_data:
        if "ип" in text_lower or "тихонов а" in text_lower:
            response = f"""📋 *Реквизиты ИП Тихонов А.В.*

*ИНН:* `{IP_TIKHONOV['inn']}`
*ОГРНИП:* `{IP_TIKHONOV['ogrnip']}`
*Адрес:* {IP_TIKHONOV['address']}

*Банк:* {IP_TIKHONOV['bank']}
*Р/с:* `{IP_TIKHONOV['rs']}`
*К/с:* `{IP_TIKHONOV['ks']}`
*БИК:* `{IP_TIKHONOV['bik']}`
"""
            await message.reply_text(response, parse_mode=ParseMode.MARKDOWN)
            return
        elif "стм" in text_lower or "ооо" in text_lower or text_lower.strip() == "реквизиты":
            await cmd_requisites(Update(0, message=message), context)
            return
    
    # Генерируем ответ через Claude
    response, model_used = await generate_response(text, file_data)
    
    # Определяем, нужен ли DOCX
    need_docx = any(word in text_lower for word in ["docx", "файл", "документ", "word", "ворд", "бланк"])
    need_docx = need_docx or (
        any(word in text_lower for word in ["составь", "напиши", "подготовь", "создай"]) and 
        any(word in text_lower for word in ["договор", "письмо", "претензи", "приказ", "иск", "акт", "заявлени"])
    )
    
    # Отправляем ответ
    if len(response) > 4000:
        if need_docx:
            docx_buffer = create_docx_on_letterhead(response)
            await message.reply_document(
                document=docx_buffer,
                filename=f"STM_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                caption="📄 Документ на фирменном бланке ООО «СТМ»"
            )
        else:
            for i in range(0, len(response), 4000):
                await message.reply_text(response[i:i+4000])
    else:
        try:
            await message.reply_text(response, parse_mode=ParseMode.MARKDOWN)
        except Exception:
            await message.reply_text(response)
        
        # Дополнительно отправляем DOCX если просили документ
        if need_docx and len(response) > 200:
            docx_buffer = create_docx_on_letterhead(response)
            await message.reply_document(
                document=docx_buffer,
                filename=f"STM_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                caption="📄 Документ на фирменном бланке"
            )
    
    # Логируем
    model_label = "💰 Sonnet" if model_used == MODEL_EXPENSIVE else "💚 Haiku"
    logger.info(f"Ответ: {model_label}, {len(response)} символов")


# ==================== ЗАПУСК ====================

def main():
    """Точка входа"""
    if not TELEGRAM_TOKEN:
        raise ValueError("TELEGRAM_TOKEN не установлен!")
    if not ANTHROPIC_API_KEY:
        raise ValueError("ANTHROPIC_API_KEY не установлен!")
    
    # Создаём папку assets если нет
    ASSETS_DIR.mkdir(exist_ok=True)
    
    # Создаём приложение
    app = Application.builder().token(TELEGRAM_TOKEN).build()
    
    # Регистрируем обработчики
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("help", cmd_help))
    app.add_handler(CommandHandler("requisites", cmd_requisites))
    
    # Обработчик всех сообщений (текст, фото, документы)
    app.add_handler(MessageHandler(
        filters.TEXT | filters.PHOTO | filters.Document.ALL,
        handle_message
    ))
    
    # Запускаем
    logger.info("🚀 Фемида v2.0 запущена!")
    logger.info(f"   Директор: @{DIRECTOR_USERNAME} (ID: {DIRECTOR_ID})")
    logger.info(f"   Группа: {GROUP_ID}")
    logger.info(f"   Модели: {MODEL_CHEAP} / {MODEL_EXPENSIVE}")
    logger.info(f"   Логотип: {'✓' if LOGO_PATH.exists() else '✗'}")
    
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
