"""Создание шаблонов бланков для каждой компании"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

TEMPLATES_DIR = Path(__file__).parent / "templates"
ASSETS_DIR = Path(__file__).parent / "assets"
LOGO_PATH = ASSETS_DIR / "logo.png"

def create_stm_letterhead():
    """Бланк ООО СТМ"""
    doc = Document()
    
    # Поля
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
    
    # Шапка с логотипом
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    left_cell = header_table.cell(0, 0)
    left_cell.width = Cm(12)
    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = p.add_run("ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ «СТМ»")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 176, 243)  # STM Blue
    
    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run("Россия, 197375, Санкт-Петербург, ул. Маршала Новикова д.42, Литер А\n")
    run2.font.size = Pt(9)
    run3 = p2.add_run("ИНН 7813568956 · КПП 781401001 · ОГРН 1137847312866")
    run3.font.size = Pt(9)
    
    # Логотип справа
    right_cell = header_table.cell(0, 1)
    right_cell.width = Cm(4)
    p_logo = right_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO_PATH), width=Cm(3))
    
    # Линия
    doc.add_paragraph("─" * 70)
    
    # Место для контента
    doc.add_paragraph()  # пустая строка
    
    doc.save(TEMPLATES_DIR / "blank_ooo_stm.docx")
    print("✓ blank_ooo_stm.docx")

def create_ip_tikhonov_letterhead():
    """Бланк ИП Тихонов А.В."""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
    
    # Шапка
    header_table = doc.add_table(rows=1, cols=2)
    left_cell = header_table.cell(0, 0)
    p = left_cell.paragraphs[0]
    
    run = p.add_run("ИНДИВИДУАЛЬНЫЙ ПРЕДПРИНИМАТЕЛЬ\nТИХОНОВ АЛЕКСАНДР ВИКТОРОВИЧ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 176, 243)
    
    p2 = left_cell.add_paragraph()
    run2 = p2.add_run("197375, Санкт-Петербург, ул. Репищева д.17, корп.1, кв.28\n")
    run2.font.size = Pt(9)
    run3 = p2.add_run("ИНН 781428127765 · ОГРНИП 319784700268498")
    run3.font.size = Pt(9)
    
    right_cell = header_table.cell(0, 1)
    p_logo = right_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO_PATH), width=Cm(3))
    
    doc.add_paragraph("─" * 70)
    doc.add_paragraph()
    
    doc.save(TEMPLATES_DIR / "blank_ip_tikhonov.docx")
    print("✓ blank_ip_tikhonov.docx")

def create_ip_trifonov_letterhead():
    """Бланк ИП Трифонов А.В."""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
    
    header_table = doc.add_table(rows=1, cols=2)
    left_cell = header_table.cell(0, 0)
    p = left_cell.paragraphs[0]
    
    run = p.add_run("ИНДИВИДУАЛЬНЫЙ ПРЕДПРИНИМАТЕЛЬ\nТРИФОНОВ АЛЕКСЕЙ ВИКТОРОВИЧ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 176, 243)
    
    p2 = left_cell.add_paragraph()
    run2 = p2.add_run("Факт: 197375, Санкт-Петербург, ул. Маршала Новикова д.42\n")
    run2.font.size = Pt(9)
    run3 = p2.add_run("ИНН 583700333358 · ОГРНИП 325580000071117")
    run3.font.size = Pt(9)
    
    right_cell = header_table.cell(0, 1)
    p_logo = right_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO_PATH), width=Cm(3))
    
    doc.add_paragraph("─" * 70)
    doc.add_paragraph()
    
    doc.save(TEMPLATES_DIR / "blank_ip_trifonov.docx")
    print("✓ blank_ip_trifonov.docx")

if __name__ == "__main__":
    TEMPLATES_DIR.mkdir(exist_ok=True)
    create_stm_letterhead()
    create_ip_tikhonov_letterhead()
    create_ip_trifonov_letterhead()
    print("\n✅ Все шаблоны созданы в templates/")
