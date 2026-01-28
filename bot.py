"""
Фемида v2.1 — Юридический ассистент ООО "СТМ"
+ Память 30 сообщений
+ Реакция на reply в группе
"""

import os
import asyncio
import io
import re
import base64
import logging
from datetime import datetime
from pathlib import Path
from collections import defaultdict
from legal_prompts import (
    detect_legal_mode, get_system_prompt, safety_check, 
    needs_escalation, ESCALATION_WARNING, MODE_EMOJI, MODE_NAME_RU
)
from legal_prompts import (
    detect_legal_mode, get_system_prompt, safety_check, 
    needs_escalation, ESCALATION_WARNING, MODE_EMOJI, MODE_NAME_RU
)

from telegram import Update, Chat, Message
from telegram.ext import Application, CommandHandler, MessageHandler, ContextTypes, filters
from telegram.constants import ParseMode, ChatAction

import anthropic
# RAG для юридической базы
import sys
sys.path.insert(0, "/opt/stm-legal-rag")
try:
    from rag_engine import get_rag
    legal_rag = get_rag()
    RAG_ENABLED = True
except Exception as e:
    legal_rag = None
    RAG_ENABLED = False
    print(f"RAG недоступен: {e}")
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

MODEL_CHEAP = "claude-3-haiku-20240307"
MODEL_EXPENSIVE = "claude-sonnet-4-20250514"
MAX_TOKENS_CHEAP = 2048
MAX_TOKENS_EXPENSIVE = 4096

DIRECTOR_USERNAME = "eugenio_spb"
DIRECTOR_ID = 1676748258
GROUP_ID = int(os.getenv("GROUP_ID", "-1003639268911"))
TRIGGERS = ["фемида,", "фемида ", "феми,", "феми ", "фем,", "фем "]
MEMORY_LIMIT = 30

ASSETS_DIR = Path(__file__).parent / "assets"
LOGO_PATH = ASSETS_DIR / "logo.png"

conversation_history = defaultdict(list)
# Кеш для media_group (несколько файлов в одном сообщении)
media_group_cache = {}
media_group_timers = {}
# Кеш последней media_group для reply (по chat_id, храним 1 час)
media_group_files_cache = {}  # {chat_id: {"files": [...], "time": datetime}}


COMPANY = {
    "full_name": "Общество с ограниченной ответственностью «СТМ»",
    "short_name": "ООО «СТМ»",
    "inn": "7813568956", "kpp": "781401001", "ogrn": "1137847312866",
    "address": "197375, Санкт-Петербург, ул. Маршала Новикова д.42, Литер А, Помещение ПИБ №1-Н-113",
    "bank": "АО «ТИНЬКОФФ БАНК»", "bik": "044525974",
    "rs": "40702810810000134609", "ks": "30101810145250000974",
    "director": "Тихонов Евгений Викторович", "director_short": "Тихонов Е.В.",
    "director_position": "Генеральный директор",
    "phone": "+7 812 603 78 71", "email": "stm.laser@gmail.com",
}

IP_TIKHONOV = {
    "full_name": "ИП Тихонов Александр Викторович", "short_name": "ИП Тихонов А.В.",
    "inn": "781428127765", "ogrnip": "319784700268498",
    "address": "197375, Санкт-Петербург, ул. Репищева д.17, корп.1, кв.28",
    "bank": "АО «ТИНЬКОФФ БАНК»", "bik": "044525974",
    "rs": "40802810400001208048", "ks": "30101810145250000974",
}

client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

EXPENSIVE_PATTERNS = [
    r"(составь|напиши|подготовь|создай|сделай).*(договор|письмо|претензи|приказ|иск|заявлени|акт)",
    r"(проанализируй|проверь|изучи|оцени).*(договор|документ|контракт)",
    r"(разработай|предложи).*(стратеги|план|схем)",
]

def is_expensive_request(text: str, has_file: bool = False) -> bool:
    if has_file:
        return True
    text_lower = text.lower()
    for pattern in EXPENSIVE_PATTERNS:
        if re.search(pattern, text_lower):
            return True
    return len(text) > 500

def get_model_for_request(text: str, has_file: bool = False) -> tuple:
    if is_expensive_request(text, has_file):
        return MODEL_EXPENSIVE, MAX_TOKENS_EXPENSIVE
    return MODEL_CHEAP, MAX_TOKENS_CHEAP

def is_director(user_id: int, username: str = None) -> bool:
    return user_id == DIRECTOR_ID or (username and username.lower() == DIRECTOR_USERNAME.lower())

def has_trigger(text: str) -> tuple:
    text_lower = text.lower()
    for trigger in TRIGGERS:
        if text_lower.startswith(trigger):
            return True, text[len(trigger):].strip()
    return False, text

async def download_file(bot, file_id: str) -> bytes:
    file = await bot.get_file(file_id)
    buffer = io.BytesIO()
    await file.download_to_memory(buffer)
    buffer.seek(0)
    return buffer.read()

async def process_document(bot, document) -> tuple:
    mime_type = document.mime_type or "application/octet-stream"
    file_data = await download_file(bot, document.file_id)
    base64_data = base64.standard_b64encode(file_data).decode("utf-8")
    if mime_type == "application/pdf":
        return base64_data, "application/pdf"
    elif mime_type.startswith("image/"):
        return base64_data, mime_type
    try:
        return file_data.decode("utf-8"), "text"
    except:
        return base64_data, mime_type

async def process_photo(bot, photo) -> tuple:
    file_data = await download_file(bot, photo.file_id)
    return base64.standard_b64encode(file_data).decode("utf-8"), "image/jpeg"

def get_current_date_ru() -> str:
    months = {1:"января",2:"февраля",3:"марта",4:"апреля",5:"мая",6:"июня",7:"июля",8:"августа",9:"сентября",10:"октября",11:"ноября",12:"декабря"}
    now = datetime.now()
    return f"{now.day} {months[now.month]} {now.year} г."

def build_system_prompt(query: str = None) -> str:
    base = f"""Ты — юридический ассистент "Фемида" компании {COMPANY['short_name']}.

ЗАДАЧИ: Составление документов, анализ договоров, консультации по ГК/ТК/НК РФ.

РЕКВИЗИТЫ ООО «СТМ»: ИНН {COMPANY['inn']}, КПП {COMPANY['kpp']}, ОГРН {COMPANY['ogrn']}
Адрес: {COMPANY['address']}
Р/с: {COMPANY['rs']}, Банк: {COMPANY['bank']}, БИК: {COMPANY['bik']}
Директор: {COMPANY['director']}

РЕКВИЗИТЫ ИП Тихонов А.В.: ИНН {IP_TIKHONOV['inn']}, ОГРНИП {IP_TIKHONOV['ogrnip']}, Р/с: {IP_TIKHONOV['rs']}

СЕГОДНЯ: {get_current_date_ru()}

Обращайся на "вы" или "Евгений". По умолчанию документы от ООО СТМ."""

    if RAG_ENABLED and legal_rag and query:
        try:
            legal_context = legal_rag.get_context_for_query(query, max_chars=2500)
            if legal_context:
                base += "\n\n📚 РЕЛЕВАНТНЫЕ СТАТЬИ ЗАКОНОДАТЕЛЬСТВА:\n" + legal_context
        except:
            pass
    return base

def add_to_memory(chat_id: int, role: str, content: str):
    conversation_history[chat_id].append({"role": role, "content": content})
    if len(conversation_history[chat_id]) > MEMORY_LIMIT:
        conversation_history[chat_id] = conversation_history[chat_id][-MEMORY_LIMIT:]

def get_memory(chat_id: int) -> list:
    return conversation_history[chat_id].copy()

def clear_memory(chat_id: int):
    conversation_history[chat_id] = []

async def generate_response(chat_id: int, text: str, file_data: tuple = None) -> tuple:
    has_file = file_data is not None
    model, max_tokens = get_model_for_request(text, has_file)
    logger.info(f"Запрос → модель: {model}, файл: {has_file}, режим: {legal_mode}")
    
    try:
        if file_data and file_data[0]:
            base64_data, media_type = file_data
            if media_type == "text":
                current_content = [{"type": "text", "text": f"Файл:\n{base64_data}\n\nЗапрос: {text}"}]
            elif media_type == "application/pdf":
                current_content = [
                    {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": base64_data}},
                    {"type": "text", "text": text or "Проанализируй этот документ."}
                ]
            else:
                current_content = [
                    {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": base64_data}},
                    {"type": "text", "text": text or "Что на этом документе?"}
                ]
        else:
            current_content = text
        
        messages = get_memory(chat_id)
        messages.append({"role": "user", "content": current_content})
        
        message = client.messages.create(model=model, max_tokens=max_tokens, system=build_system_prompt(text), messages=messages)
        response_text = message.content[0].text
        
        add_to_memory(chat_id, "user", text)
        add_to_memory(chat_id, "assistant", response_text)
        
        logger.info(f"Токены: in={message.usage.input_tokens}, out={message.usage.output_tokens}, память: {len(get_memory(chat_id))}")
        return response_text, model
    except Exception as e:
        logger.error(f"Ошибка Claude: {e}")
        return f"⚠️ Ошибка: {e}", model

def create_docx_on_letterhead(content: str) -> io.BytesIO:
    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    
    header = doc.sections[0].header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.columns[0].width = Inches(1.2)
    header_table.columns[1].width = Inches(5.3)
    
    logo_cell = header_table.cell(0, 0)
    if LOGO_PATH.exists():
        logo_para = logo_cell.paragraphs[0]
        logo_run = logo_para.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Inches(1))
    
    text_cell = header_table.cell(0, 1)
    name_para = text_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = name_para.add_run("ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ")
    run1.font.bold = True
    run1.font.size = Pt(11)
    
    p2 = text_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("«СТМ»")
    run2.font.bold = True
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0, 112, 192)
    
    p3 = text_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"{COMPANY['address']}\nИНН {COMPANY['inn']} · КПП {COMPANY['kpp']} · ОГРН {COMPANY['ogrn']}")
    run3.font.size = Pt(8)
    
    line = header.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = line.add_run("─" * 85)
    lr.font.size = Pt(8)
    lr.font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph()
    for para_text in content.split('\n'):
        if para_text.strip():
            p = doc.add_paragraph()
            stripped = para_text.strip()
            if stripped.isupper() or any(stripped.startswith(x) for x in ['ДОГОВОР','ПРИКАЗ','ПРЕТЕНЗИЯ','АКТ','ПИСЬМО']):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.bold = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                run = p.add_run(stripped)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run(f"{COMPANY['director_position']}                    _____________    {COMPANY['director_short']}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    if not is_director(user.id, user.username):
        await update.message.reply_text("⚖️ Фемида доступна только руководству ООО «СТМ».")
        return
    await update.message.reply_text(
        "⚖️ *Фемида v2.1* — юридический ассистент ООО «СТМ»\n\n"
        "• Составление договоров, писем, претензий\n"
        "• Анализ документов (PDF, фото)\n"
        "• Юридические консультации\n\n"
        f"_Память: {MEMORY_LIMIT} сообщений_\n"
        "_В группе: Фемида, ... или ответ на моё сообщение_\n\n"
        "/clear — очистить память",
        parse_mode=ParseMode.MARKDOWN
    )

async def cmd_clear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    clear_memory(chat_id)
    await update.message.reply_text("🧹 Память очищена.")

async def cmd_requisites(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        f"📋 *Реквизиты ООО «СТМ»*\n\n"
        f"ИНН: `{COMPANY['inn']}`\nКПП: `{COMPANY['kpp']}`\nОГРН: `{COMPANY['ogrn']}`\n"
        f"Адрес: {COMPANY['address']}\n\n"
        f"Банк: {COMPANY['bank']}\nР/с: `{COMPANY['rs']}`\nБИК: `{COMPANY['bik']}`\n\n"
        f"Директор: {COMPANY['director']}",
        parse_mode=ParseMode.MARKDOWN
    )

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    message = update.message
    if not message:
        return
    
    user = update.effective_user
    chat = update.effective_chat
    bot_id = context.bot.id
    
    # === ОБРАБОТКА MEDIA_GROUP (несколько файлов) ===
    if message.media_group_id:
        mg_id = message.media_group_id
        
        # Инициализируем кеш
        if mg_id not in media_group_cache:
            media_group_cache[mg_id] = {"files": [], "text": "", "message": message, "user": user, "chat": chat}
        
        # Собираем файл
        if message.document:
            fd = await process_document(context.bot, message.document)
            if fd:
                media_group_cache[mg_id]["files"].append(fd)
        elif message.photo:
            fd = await process_photo(context.bot, message.photo[-1])
            if fd:
                media_group_cache[mg_id]["files"].append(fd)
        
        # Сохраняем caption
        if message.caption:
            media_group_cache[mg_id]["text"] = message.caption.strip()
        
        # Отменяем предыдущий таймер
        if mg_id in media_group_timers:
            media_group_timers[mg_id].cancel()
        
        # Таймер на обработку
        async def process_mg():
            await asyncio.sleep(1.5)
            if mg_id in media_group_cache:
                data = media_group_cache.pop(mg_id)
                media_group_timers.pop(mg_id, None)
                files_list = data["files"]
                txt = data["text"] or f"Проанализируй эти {len(files_list)} документов."
                msg = data["message"]
                usr = data["user"]
                cht = data["chat"]
                
                if cht.type in [Chat.GROUP, Chat.SUPERGROUP]:
                    if not is_director(usr.id, usr.username):
                        return
                    has_trig, clean_txt = has_trigger(txt)
                    if has_trig:
                        await process_request_multi(msg, clean_txt, files_list, context)
                elif cht.type == Chat.PRIVATE:
                    if is_director(usr.id, usr.username):
                        await process_request_multi(msg, txt, files_list, context)
        
        task = asyncio.create_task(process_mg())
        media_group_timers[mg_id] = task
        return
    
    # === ОБЫЧНАЯ ОБРАБОТКА (один файл) ===
    text = message.text or message.caption or ""
    text = text.strip()
    
    file_data = None
    if message.document:
        file_data = await process_document(context.bot, message.document)
    elif message.photo:
        file_data = await process_photo(context.bot, message.photo[-1])
    
    if file_data and not text:
        text = "Проанализируй этот документ."
    
    if not text and not file_data:
        return
    
    if chat.type == Chat.PRIVATE:
        if not is_director(user.id, user.username):
            await message.reply_text("⚖️ Фемида доступна только руководству ООО «СТМ».")
            return
        await process_request(message, text, file_data, context)
        return
    
    if chat.type in [Chat.GROUP, Chat.SUPERGROUP]:
        if not is_director(user.id, user.username):
            return
        
        has_trig, clean_text = has_trigger(text)
        is_reply_to_bot = message.reply_to_message and message.reply_to_message.from_user and message.reply_to_message.from_user.id == bot_id
        

        # Берём файлы из reply если в текущем сообщении нет
        if not file_data and message.reply_to_message:
            reply_msg = message.reply_to_message
            reply_id = reply_msg.message_id
            
            # Проверяем кеш media_group (по chat_id)
            if chat.id in media_group_files_cache:
                cached = media_group_files_cache[chat.id]
                if (datetime.now() - cached["time"]).seconds < 3600:
                    files_list = cached["files"]
                    has_trig, clean_text = has_trigger(text)
                    if has_trig:
                        await process_request_multi(message, clean_text, files_list, context)
                        return
            
            # Обычная обработка одного файла из reply
            if reply_msg.document:
                file_data = await process_document(context.bot, reply_msg.document)
            elif reply_msg.photo:
                file_data = await process_photo(context.bot, reply_msg.photo[-1])
        if has_trig:
            await process_request(message, clean_text, file_data, context)
        elif is_reply_to_bot:
            await process_request(message, text, file_data, context)


async def process_request_multi(message: Message, text: str, files_list: list, context: ContextTypes.DEFAULT_TYPE):
    """Обработка запроса с несколькими файлами"""
    chat_id = message.chat_id
    
    # Сохраняем в кеш для будущих reply (по chat_id)
    media_group_files_cache[chat_id] = {
        "files": files_list,
        "time": datetime.now()
    }
    # Чистим старые записи (>1 часа)
    old_ids = [k for k, v in media_group_files_cache.items() 
               if (datetime.now() - v["time"]).seconds > 3600]
    for k in old_ids:
        media_group_files_cache.pop(k, None)
    await context.bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)
    
    # Собираем все файлы в один запрос
    combined_content = []
    for i, file_data in enumerate(files_list, 1):
        base64_data, media_type = file_data
        combined_content.append({
            "type": "document" if media_type == "application/pdf" else "image",
            "source": {"type": "base64", "media_type": media_type, "data": base64_data}
        })
    
    # Формируем сообщение для Claude
    messages_content = combined_content + [{"type": "text", "text": text}]
    
    model = MODEL_EXPENSIVE  # Всегда Sonnet для мультифайлов
    
    # RAG контекст
    legal_context = ""
    if RAG_ENABLED and legal_rag:
        legal_context = legal_rag.get_context_for_query(text)
    
    system, legal_mode, escalation_flag = build_system_prompt(text, has_file)
    if legal_context:
        system += f"\n\nПРАВОВАЯ БАЗА:\n{legal_context}"
    
    try:
        response = client.messages.create(
            model=model,
            max_tokens=MAX_TOKENS_EXPENSIVE,
            system=system,
            messages=[{"role": "user", "content": messages_content}]
        )
        result = response.content[0].text
        logger.info(f"Мультифайл: {len(files_list)} файлов, токены: in={response.usage.input_tokens}, out={response.usage.output_tokens}")
    except Exception as e:
        logger.error(f"Ошибка Claude: {e}")
        result = f"Ошибка обработки: {e}"
    
    # Отправляем ответ
    text_lower = text.lower()
    need_docx = any(phrase in text_lower for phrase in ["на бланке", "создай ответ", "создай письмо", "подготовь ответ"])
    
    if len(result) > 4000:
        for i in range(0, len(result), 4000):
            await message.reply_text(result[i:i+4000])
    else:
        try:
            await message.reply_text(result, parse_mode=ParseMode.MARKDOWN)
        except:
            await message.reply_text(result)
    
    if need_docx:
        from companies import find_company
        company_key, company_data = find_company(text_lower)
        if company_key:
            docx_buffer = create_docx_on_letterhead(result)
            await message.reply_document(
                document=docx_buffer,
                filename=f"STM_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                caption=f"📄 На бланке {company_data['short_name']}"
            )

async def process_request(message: Message, text: str, file_data: tuple, context: ContextTypes.DEFAULT_TYPE):
    chat_id = message.chat_id
    await context.bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)
    
    text_lower = text.lower()
    if "реквизиты" in text_lower and not file_data:
        if "ип" in text_lower:
            await message.reply_text(
                f"📋 *Реквизиты ИП Тихонов А.В.*\n\nИНН: `{IP_TIKHONOV['inn']}`\nОГРНИП: `{IP_TIKHONOV['ogrnip']}`\n"
                f"Адрес: {IP_TIKHONOV['address']}\n\nБанк: {IP_TIKHONOV['bank']}\nР/с: `{IP_TIKHONOV['rs']}`\nБИК: `{IP_TIKHONOV['bik']}`",
                parse_mode=ParseMode.MARKDOWN
            )
            return
        elif "стм" in text_lower or "ооо" in text_lower or text_lower.strip() == "реквизиты":
            await cmd_requisites(Update(0, message=message), context)
            return
    
    response, model_used = await generate_response(chat_id, text, file_data)
    

    # DOCX только по явному запросу: "создай на бланке ИП/ООО/Трифонова"
    need_docx = False
    company_key = None
    if any(phrase in text_lower for phrase in ["на бланке", "на бланк ", "создай ответ", "создай письмо", "создай претензию", "подготовь ответ", "подготовь письмо"]):
        from companies import find_company
        company_key, company_data = find_company(text_lower)
        if company_key:
            need_docx = True
    
    if len(response) > 4000:
        if need_docx:
            docx_buffer = create_docx_on_letterhead(response)
            await message.reply_document(document=docx_buffer, filename=f"STM_{datetime.now().strftime('%Y%m%d_%H%M')}.docx", caption="📄 Документ на бланке СТМ")
        else:
            for i in range(0, len(response), 4000):
                await message.reply_text(response[i:i+4000])
    else:
        try:
            await message.reply_text(response, parse_mode=ParseMode.MARKDOWN)
        except:
            await message.reply_text(response)
        if need_docx and len(response) > 200:
            docx_buffer = create_docx_on_letterhead(response)
            await message.reply_document(document=docx_buffer, filename=f"STM_{datetime.now().strftime('%Y%m%d_%H%M')}.docx", caption="📄 На бланке")
    
    logger.info(f"{'💰 Sonnet' if model_used == MODEL_EXPENSIVE else '💚 Haiku'}, {len(response)} симв.")

def main():
    if not TELEGRAM_TOKEN or not ANTHROPIC_API_KEY:
        raise ValueError("Токены не установлены!")
    ASSETS_DIR.mkdir(exist_ok=True)
    app = Application.builder().token(TELEGRAM_TOKEN).build()
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("help", cmd_start))
    app.add_handler(CommandHandler("clear", cmd_clear))
    app.add_handler(CommandHandler("requisites", cmd_requisites))
    app.add_handler(MessageHandler(filters.TEXT | filters.PHOTO | filters.Document.ALL, handle_message))
    logger.info(f"🚀 Фемида v2.1 | Память: {MEMORY_LIMIT} | Логотип: {'✓' if LOGO_PATH.exists() else '✗'}")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
